"""
Corpus export helper — Phase 0 + Phase 2 fixtures.

Phase 0 fixtures:
  1. Loads HuggingFace Arrow dataset → 5 text samples
  2. Synthetic Indian CA resume → .txt / .pdf / .docx / .jpg
  3. Empty test workbook

Phase 2 fixtures (routing tests):
  4. sample_scanned.pdf       — image-only PDF (no text layer)
  5. sample_garbled.pdf       — PDF with garbage Unicode text layer
  6. sample_multicolumn.pdf   — two-column layout PDF
  7. sample_screenshot.jpg    — low-DPI phone-screenshot-style JPG
  8. sample_legacy.doc        — .docx content renamed to .doc (stub)
  9. sample_corrupt.pdf       — zero-byte invalid PDF

Run:  python test_corpus/export_corpus.py
Output: test_corpus/files/*  +  test_corpus/test_output.xlsx
"""
from __future__ import annotations

import sys
from pathlib import Path

# Ensure repo root is on path when run directly
ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(ROOT))

CORPUS_DIR = ROOT / "test_corpus" / "resumeatlas_sample"
FILES_DIR = ROOT / "test_corpus" / "files"
OUTPUT_DIR = ROOT / "output_data"
WORKBOOK_PATH = ROOT / "test_corpus" / "test_output.xlsx"


# ---------------------------------------------------------------------------
# Synthetic Indian CA resume content
# ---------------------------------------------------------------------------

CA_RESUME_TEXT = """\
RAHUL MEHTA
CA (Final) | ICAI Membership No. 123456
rahul.mehta@example.com | +91-98765-43210
Mumbai, Maharashtra, India
LinkedIn: linkedin.com/in/rahulmehta-ca

SUMMARY
Chartered Accountant with 3 years of articleship at Big-4 firm. Specialisation in
statutory audit, direct taxation, and GST advisory. Cleared CA Final in November 2023
attempt (both groups). Currently seeking a role in financial advisory or audit.

PROFESSIONAL QUALIFICATIONS
- CA Final | ICAI | Cleared | November 2023 | 2 attempts
- CA Intermediate | ICAI | Cleared | May 2021 | 1 attempt
- CA Foundation | ICAI | Cleared | December 2019 | 1 attempt

ICAI Membership Number: 123456 (enrolled January 2024)

ARTICLESHIP / INTERNSHIP
Firm: Deloitte Haskins & Sells LLP, Mumbai
Role: Article Assistant
Period: January 2021 – December 2023 (36 months)
Areas: Statutory Audit, Tax Audit, GST Compliance, Direct Taxation,
       Transfer Pricing, Internal Audit, Company Law

EDUCATION
B.Com (Hons) | University of Mumbai | 2017 – 2020 | 78% | Completed
HSC (Commerce) | Maharashtra State Board | 2015 – 2017 | 85% | Completed
SSC | Maharashtra State Board | 2015 | 91% | Completed

WORK EXPERIENCE
Organisation: Deloitte Haskins & Sells LLP
Designation: Article Assistant
Type: Intern
Period: January 2021 – December 2023
Location: Mumbai, Maharashtra
Responsibilities:
  - Conducted statutory audits for listed and unlisted companies
  - Prepared tax audit reports under Section 44AB of Income Tax Act
  - Assisted in GST reconciliation and filing (GSTR-1, GSTR-3B, GSTR-9)
  - Reviewed internal controls and prepared audit observations
  - Assisted in transfer pricing documentation for multinational clients

SKILLS
Technical: Tally ERP 9, SAP FICO, MS Excel (Advanced), IRIS (MCA filing), GST portal
Domain: Statutory Audit, Tax Audit, GST, Direct Tax, Transfer Pricing, IFRS/Ind AS
Soft Skills: Analytical Thinking, Team Collaboration, Report Writing

ACHIEVEMENTS
- All India Rank 42 in CA Intermediate (May 2021)
- Represented college at inter-college commerce fest (2019)

LANGUAGES
English (Fluent), Hindi (Native), Marathi (Native)
"""

# ---------------------------------------------------------------------------
# Export functions
# ---------------------------------------------------------------------------

def export_text_samples(n: int = 5) -> list[Path]:
    """Load Arrow dataset and write first n samples as .txt files."""
    from datasets import load_from_disk

    print(f"Loading Arrow dataset from {CORPUS_DIR} …")
    ds = load_from_disk(str(CORPUS_DIR))
    total = len(ds)
    print(f"  {total} rows found. Exporting {min(n, total)} samples.")

    exported = []
    for i in range(min(n, total)):
        category = ds[i]["Category"].replace(" ", "_").replace("/", "-")
        text = ds[i]["Text"]
        fname = FILES_DIR / f"sample_{i+1:02d}_{category}.txt"
        fname.write_text(text, encoding="utf-8")
        exported.append(fname)
        print(f"  Exported: {fname.name}  ({len(text)} chars)")
    return exported


def export_ca_text() -> Path:
    """Write the synthetic Indian CA resume as a .txt file."""
    path = FILES_DIR / "sample_ca_india_synthetic.txt"
    path.write_text(CA_RESUME_TEXT, encoding="utf-8")
    print(f"  Exported: {path.name}")
    return path


def export_ca_pdf() -> Path:
    """Generate a minimal single-page PDF from the CA resume text."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    path = FILES_DIR / "sample_ca_india_synthetic.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=A4,
                            leftMargin=50, rightMargin=50,
                            topMargin=50, bottomMargin=50)
    styles = getSampleStyleSheet()
    story = []
    for line in CA_RESUME_TEXT.splitlines():
        if line.strip():
            story.append(Paragraph(line.replace("&", "&amp;"), styles["Normal"]))
        else:
            story.append(Spacer(1, 6))
    doc.build(story)
    print(f"  Exported: {path.name}")
    return path


def export_ca_docx() -> Path:
    """Generate a minimal DOCX from the CA resume text."""
    from docx import Document

    path = FILES_DIR / "sample_ca_india_synthetic.docx"
    doc = Document()
    for line in CA_RESUME_TEXT.splitlines():
        doc.add_paragraph(line)
    doc.save(str(path))
    print(f"  Exported: {path.name}")
    return path


def export_ca_jpg() -> Path:
    """Generate a minimal JPG image rendering of the CA resume (first page)."""
    from PIL import Image, ImageDraw, ImageFont

    path = FILES_DIR / "sample_ca_india_synthetic.jpg"
    img = Image.new("RGB", (794, 1123), color="white")  # A4 at 96 DPI
    draw = ImageDraw.Draw(img)

    try:
        font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 11)
        font_bold = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 13)
    except Exception:
        font = ImageFont.load_default()
        font_bold = font

    y = 40
    for line in CA_RESUME_TEXT.splitlines()[:60]:  # first 60 lines fit on one page
        f = font_bold if line.isupper() or line.startswith("RAHUL") else font
        draw.text((40, y), line, fill="black", font=f)
        y += 16
        if y > 1080:
            break

    img.save(str(path), "JPEG", quality=85)
    print(f"  Exported: {path.name}")
    return path


# ---------------------------------------------------------------------------
# Phase 2 fixtures — routing test corpus
# ---------------------------------------------------------------------------

_SHORT_RESUME = """\
John Smith
john.smith@example.com | +1-555-0100
Software Engineer, Acme Corp, 2020-2023
B.Sc. Computer Science, State University, 2016-2020
Skills: Python, Java, SQL
"""


def export_scanned_pdf() -> Path:
    """
    PDF where every page is a rasterized image — no text layer.
    Created by rendering text to a Pillow image and embedding it in a
    reportlab PDF that contains only the image (no text drawing commands).
    """
    from PIL import Image, ImageDraw, ImageFont
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate
    from reportlab.platypus import Image as RLImage
    import io

    # Step 1: render text to a Pillow image
    img = Image.new("RGB", (794, 400), "white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 14)
    except Exception:
        font = ImageFont.load_default()
    y = 30
    for line in _SHORT_RESUME.splitlines():
        draw.text((30, y), line, fill="black", font=font)
        y += 22

    # Step 2: save PIL image to bytes
    img_bytes = io.BytesIO()
    img.save(img_bytes, "PNG")
    img_bytes.seek(0)

    # Step 3: embed image-only in reportlab PDF (no text objects)
    path = FILES_DIR / "sample_scanned.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=A4,
                            leftMargin=20, rightMargin=20, topMargin=20, bottomMargin=20)
    rl_img = RLImage(img_bytes, width=500, height=250)
    doc.build([rl_img])
    print(f"  Exported: {path.name}  (image-only PDF, no text layer)")
    return path


def export_garbled_pdf() -> Path:
    """
    PDF with a text layer that is garbage Unicode (high junk-char ratio,
    no real words, no spaces). Simulates a broken font-encoding export.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    import random, string

    path = FILES_DIR / "sample_garbled.pdf"
    c = canvas.Canvas(str(path), pagesize=A4)
    c.setFont("Helvetica", 10)

    # Generate garbled content: random mix of rare Unicode ranges + no spaces
    random.seed(42)
    garbled_lines = []
    for _ in range(40):
        # Mix of Combining Diacritical Marks (U+0300-U+036F) range chars and
        # random non-Latin characters — high junk ratio, no real words
        line = "".join(
            chr(random.choice(
                list(range(0x0300, 0x0370)) +   # combining marks
                list(range(0x0590, 0x05FF)) +   # Hebrew (not readable as resume)
                list(range(0x0600, 0x060F))     # Arabic control chars
            ))
            for _ in range(random.randint(30, 60))
        )
        garbled_lines.append(line)

    y = 800
    for line in garbled_lines:
        try:
            c.drawString(30, y, line[:40])  # reportlab may skip unprintable chars
        except Exception:
            c.drawString(30, y, "XXXX" * 10)  # fallback: clearly junk
        y -= 18

    c.save()
    print(f"  Exported: {path.name}  (garbled Unicode text layer)")
    return path


def export_multicolumn_pdf() -> Path:
    """
    PDF with two-column layout where linear text extraction scrambles reading order.
    Uses reportlab Frames to place text in left and right columns.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import BaseDocTemplate, Frame, PageTemplate, Paragraph, Spacer

    path = FILES_DIR / "sample_multicolumn.pdf"
    W, H = A4  # 595.27 x 841.89

    # Two equal columns with a gutter
    col_w = (W - 100) / 2  # ~247 pts each
    left_frame = Frame(30, 50, col_w, H - 100, id="left")
    right_frame = Frame(W / 2 + 10, 50, col_w, H - 100, id="right")

    doc = BaseDocTemplate(str(path), pagesize=A4)
    doc.addPageTemplates([PageTemplate(id="two_col", frames=[left_frame, right_frame])])

    from reportlab.platypus import FrameBreak

    styles = getSampleStyleSheet()
    story = []
    # Left column content
    left_text = [
        "PRIYA SHARMA", "priya@example.com", "+91-99887-76655",
        "Mumbai, Maharashtra", "",
        "EXPERIENCE", "Senior Auditor", "Ernst & Young LLP",
        "Jan 2022 – Present", "Led statutory audits", "GST compliance",
        "Transfer pricing documentation", "", "EDUCATION",
        "B.Com (Hons)", "University of Mumbai", "2015-2018", "78%",
    ]
    right_text = [
        "QUALIFICATIONS", "CA Final – Cleared Nov 2022",
        "ICAI Membership: 234567", "CA Intermediate – May 2020",
        "CA Foundation – Dec 2018", "",
        "SKILLS", "Tally ERP", "SAP FICO", "MS Excel", "Direct Tax",
        "Statutory Audit", "GST Advisory", "",
        "LANGUAGES", "English (Fluent)", "Hindi (Native)",
    ]
    for line in left_text:
        if line:
            story.append(Paragraph(line.replace("&", "&amp;"), styles["Normal"]))
        else:
            story.append(Spacer(1, 12))
    # Force a break to the right frame so blocks appear at right-side x positions
    story.append(FrameBreak())
    for line in right_text:
        if line:
            story.append(Paragraph(line.replace("&", "&amp;"), styles["Normal"]))
        else:
            story.append(Spacer(1, 12))

    doc.build(story)
    print(f"  Exported: {path.name}  (two-column layout)")
    return path


def export_screenshot_jpg() -> Path:
    """
    Low-DPI (96 dpi) phone-screenshot-style JPG — simulates a photo of a resume.
    Text rendered at small size onto a white background.
    """
    from PIL import Image, ImageDraw, ImageFont

    # 96 DPI A4 ≈ 794×1123 px → use that size but low quality
    path = FILES_DIR / "sample_screenshot.jpg"
    img = Image.new("RGB", (600, 800), color=(245, 245, 240))  # slightly off-white
    draw = ImageDraw.Draw(img)

    try:
        font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 10)
        font_hdr = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 12)
    except Exception:
        font = ImageFont.load_default()
        font_hdr = font

    y = 20
    for line in _SHORT_RESUME.splitlines():
        f = font_hdr if line and line[0].isupper() and len(line) < 30 else font
        draw.text((15, y), line, fill=(20, 20, 20), font=f)
        y += 18
        if y > 780:
            break

    # Save at low quality to mimic phone photo
    img.save(str(path), "JPEG", quality=55, dpi=(96, 96))
    print(f"  Exported: {path.name}  (96dpi screenshot-style JPG)")
    return path


def export_legacy_doc() -> Path:
    """
    Legacy .doc stub: create a DOCX with python-docx, then rename to .doc.
    The file is structurally a DOCX (ZIP-based OOXML), not a real OLE .doc,
    but serves as a fixture to test the dead-letter-or-convert path.
    Without LibreOffice the router will dead-letter it with a clear reason.
    """
    from docx import Document

    path = FILES_DIR / "sample_legacy.doc"
    doc = Document()
    doc.add_heading("Legacy DOC Fixture", level=1)
    doc.add_paragraph(_SHORT_RESUME)
    doc.save(str(path))
    print(f"  Exported: {path.name}  (.docx-as-.doc stub; dead-letters without LibreOffice)")
    return path


def export_corrupt_pdf() -> Path:
    """Zero-byte file — cannot be opened by any PDF parser."""
    path = FILES_DIR / "sample_corrupt.pdf"
    path.write_bytes(b"")
    print(f"  Exported: {path.name}  (zero-byte corrupt PDF)")
    return path


def create_test_workbook() -> Path:
    """Create an empty test workbook with placeholder sheets."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws_main = wb.active
    ws_main.title = "Candidates"
    ws_main.append(["[Phase 4 will populate headers]"])

    ws_review = wb.create_sheet("Review Queue")
    ws_review.append(["[Low-confidence / failed parses land here]"])

    wb.save(str(WORKBOOK_PATH))
    print(f"  Created: {WORKBOOK_PATH.name}")
    return WORKBOOK_PATH


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    FILES_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    print("\n=== Corpus export ===")
    exported = []

    print("\n[Phase 0] Text samples from Arrow dataset")
    exported.extend(export_text_samples(n=5))

    print("\n[Phase 0] Synthetic Indian CA resume → .txt/.pdf/.docx/.jpg")
    exported.append(export_ca_text())
    exported.append(export_ca_pdf())
    exported.append(export_ca_docx())
    exported.append(export_ca_jpg())

    print("\n[Phase 2] Routing test fixtures")
    exported.append(export_scanned_pdf())
    exported.append(export_garbled_pdf())
    exported.append(export_multicolumn_pdf())
    exported.append(export_screenshot_jpg())
    exported.append(export_legacy_doc())
    exported.append(export_corrupt_pdf())

    print("\n[+] Creating test workbook")
    create_test_workbook()

    print(f"\nDone. {len(exported)} files in {FILES_DIR}")
    for f in sorted(FILES_DIR.iterdir()):
        print(f"  {f.name}  ({f.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
